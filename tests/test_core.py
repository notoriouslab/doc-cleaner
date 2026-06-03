"""Unit tests for core.py — shared conversion API (tasks 1.1–1.4)."""
import os
import re
import tempfile
import pytest
from pathlib import Path
from unittest.mock import patch

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

PUBDATE_RE = re.compile(r'^pubDate: ".*"$', re.MULTILINE)


def _strip_timestamp(text):
    return PUBDATE_RE.sub('pubDate: "<stripped>"', text)


@pytest.fixture
def sample_docx(tmp_path):
    """Minimal .docx for testing."""
    if not HAS_DOCX:
        pytest.skip("python-docx not installed")
    doc = Document()
    doc.add_paragraph("Test content for core.py.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def two_valid_docx(tmp_path):
    """Two valid .docx files."""
    if not HAS_DOCX:
        pytest.skip("python-docx not installed")
    paths = []
    for i in range(2):
        doc = Document()
        doc.add_paragraph(f"Document {i + 1}")
        p = tmp_path / f"doc{i + 1}.docx"
        doc.save(str(p))
        paths.append(str(p))
    return paths


# ── Task 1.1: convert_file returns ok and writes beside source ────────────────

class TestConvertFile:
    def test_ok_status_and_output_written(self, sample_docx, tmp_path):
        """convert_file on a valid .docx returns status=ok with a written .md."""
        out_dir = str(tmp_path / "out")
        os.makedirs(out_dir)

        from core import convert_file
        result = convert_file(sample_docx, output_dir=out_dir)

        assert result["status"] == "ok"
        assert result["output"] is not None
        assert Path(result["output"]).exists()

    def test_output_dir_none_writes_beside_source(self, sample_docx, tmp_path):
        """output_dir=None writes the .md beside the source file."""
        # Move sample into its own subdirectory so we can check parentage
        src_dir = tmp_path / "src"
        src_dir.mkdir()
        import shutil
        dest = str(src_dir / "sample.docx")
        shutil.copy(sample_docx, dest)

        from core import convert_file
        result = convert_file(dest, output_dir=None)

        assert result["status"] == "ok"
        assert Path(result["output"]).parent == src_dir


# ── Task 1.2: status mapping from the four underlying raw statuses ────────────

class TestStatusMapping:
    @pytest.mark.parametrize("raw,expected", [
        ("ok", "ok"),
        ("dry_run", "dry_run"),
        ("no_content", "skipped"),
        ("write_error", "error"),
        ("error", "error"),
    ])
    def test_status_map(self, raw, expected, tmp_path):
        """Each raw process_file status maps to the correct API status."""
        from core import _run_one

        fake_file = str(tmp_path / "fake.docx")
        Path(fake_file).write_bytes(b"dummy")

        out_path = str(tmp_path / "out.md") if raw == "ok" else None
        if out_path:
            Path(out_path).write_text("# Test", encoding="utf-8")

        with patch("core.process_file", return_value=(raw, out_path)):
            result = _run_one(fake_file, None, None, {}, str(tmp_path))

        assert result["status"] == expected

    def test_error_status_carries_message(self, tmp_path):
        """Error status captures the log message emitted by process_file."""
        from core import _run_one
        import logging

        fake_file = str(tmp_path / "fake.docx")
        Path(fake_file).write_bytes(b"dummy")

        def _fake_process_file(*args, **kwargs):
            logging.getLogger("doc-cleaner").error("something went wrong")
            return ("error", None)

        with patch("core.process_file", side_effect=_fake_process_file):
            result = _run_one(fake_file, None, None, {}, str(tmp_path))

        assert result["status"] == "error"
        assert result["error"] is not None
        assert "something went wrong" in result["error"]

    def test_error_status_fallback_message(self, tmp_path):
        """Error status uses fallback '轉換失敗' when no log message was captured."""
        from core import _run_one

        fake_file = str(tmp_path / "fake.docx")
        Path(fake_file).write_bytes(b"dummy")

        with patch("core.process_file", return_value=("error", None)):
            result = _run_one(fake_file, None, None, {}, str(tmp_path))

        assert result["status"] == "error"
        assert result["error"] == "轉換失敗"

    def test_skipped_status_has_reason(self, tmp_path):
        """Skipped status always carries a user-readable reason (not None)."""
        from core import _run_one

        fake_file = str(tmp_path / "fake.docx")
        Path(fake_file).write_bytes(b"dummy")

        with patch("core.process_file", return_value=("no_content", None)):
            result = _run_one(fake_file, None, None, {}, str(tmp_path))

        assert result["status"] == "skipped"
        assert result["error"] == "未擷取到文字內容"

    def test_parsers_logger_captured(self, tmp_path):
        """Errors from the 'parsers' logger namespace are captured in the result."""
        from core import _run_one
        import logging

        fake_file = str(tmp_path / "fake.xlsx")
        Path(fake_file).write_bytes(b"dummy")

        def _fake_process_file(*args, **kwargs):
            logging.getLogger("parsers.xlsx").error("Excel parse failed: bad zip")
            return ("error", None)

        with patch("core.process_file", side_effect=_fake_process_file):
            result = _run_one(fake_file, None, None, {}, str(tmp_path))

        assert result["error"] is not None
        assert "Excel parse failed" in result["error"]


# ── Task 1.3: batch reuses one backend and continues past errors ──────────────

class TestConvertFiles:
    def test_batch_three_files_middle_errors(self, two_valid_docx, tmp_path):
        """3-file batch: middle file errors, batch produces 3 results, only middle is error."""
        valid1, valid2 = two_valid_docx

        # A file whose path will be intercepted by the mock
        broken = str(tmp_path / "broken.docx")
        Path(broken).write_bytes(b"dummy")

        out_dir = str(tmp_path / "output")
        os.makedirs(out_dir)

        # Build an output path for the two successful files
        ok_out = str(tmp_path / "ok.md")
        Path(ok_out).write_text("# ok", encoding="utf-8")

        broken_abs = str(Path(broken).resolve())

        def _fake_process(filepath, *args, **kwargs):
            if str(Path(filepath).resolve()) == broken_abs:
                return ("error", None)
            return ("ok", ok_out)

        from core import convert_files
        with patch("core.process_file", side_effect=_fake_process):
            results = convert_files(
                [valid1, broken, valid2],
                output_resolver=lambda _: out_dir,
            )

        assert len(results) == 3
        assert results[0]["status"] == "ok"
        assert results[1]["status"] == "error"
        assert results[2]["status"] == "ok"

    def test_batch_builds_backend_once(self, two_valid_docx, tmp_path):
        """Batch calls process_file per file but create_ai_backend only once."""
        out_dir = str(tmp_path / "output")
        os.makedirs(out_dir)

        from core import convert_files
        with patch("core.create_ai_backend", wraps=__import__("cleaner").create_ai_backend) as mock_backend:
            convert_files(two_valid_docx, output_resolver=lambda _: out_dir)
        # create_ai_backend called once for the whole batch, not per-file
        assert mock_backend.call_count == 1


# ── Task 1.4: CLI output unchanged after delegation to core.py ───────────────

class TestCLIOutputUnchanged:
    def test_dry_run_status_preserved(self, tmp_path):
        """dry_run raw status maps to 'dry_run' (not 'ok') so --summary is unchanged."""
        from core import _run_one

        fake_file = str(tmp_path / "fake.docx")
        Path(fake_file).write_bytes(b"dummy")
        out_path = str(tmp_path / "fake.md")

        with patch("core.process_file", return_value=("dry_run", out_path)):
            result = _run_one(fake_file, None, None, {}, str(tmp_path))

        assert result["status"] == "dry_run"

    def test_markdown_content_identical(self, sample_docx, tmp_path):
        """
        convert_file produces the same Markdown as a direct process_file call,
        excluding the pubDate timestamp which differs by wall-clock time.
        """
        from cleaner import load_config, process_file, SCRIPT_DIR
        from core import convert_file

        # Direct process_file path
        dir_direct = str(tmp_path / "direct")
        os.makedirs(dir_direct)
        config = load_config(str(SCRIPT_DIR / "config.json"))
        _, out_direct = process_file(sample_docx, None, None, config, dir_direct)

        # core.py path (different output dir to avoid filename collision)
        dir_core = str(tmp_path / "core")
        os.makedirs(dir_core)
        result = convert_file(sample_docx, output_dir=dir_core)
        out_core = result["output"]

        content_direct = Path(out_direct).read_text(encoding="utf-8")
        content_core = Path(out_core).read_text(encoding="utf-8")

        assert _strip_timestamp(content_direct) == _strip_timestamp(content_core)
